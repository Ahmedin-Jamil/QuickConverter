import pdfplumber
import pandas as pd
import re
import hashlib
import logging
from typing import List, Dict, Any, Union
from abc import ABC, abstractmethod

class BaseParser(ABC):
    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, Any]:
        pass

    def get_file_hash(self, file_path: str) -> str:
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

class PDFParser(BaseParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Returns a hybrid payload:
        {
            "hash": "...",
            "fragments": [{"type": "table_row", "data": [...], "page": 1}, ...],
            "raw_text": "full text blob..."
        }
        """
        fragments = []
        raw_text_pages = []
        file_hash = self.get_file_hash(file_path)
        
        logging.info(f"Hybrid Extracting PDF: {file_path}")
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                # 1. Capture tables
                tables = page.extract_tables()
                for table in tables:
                    if self._is_likely_transaction_table(table):
                        for row in table:
                            fragments.append({
                                "type": "table_row", 
                                "data": row, 
                                "page_number": i+1
                            })
                
                # 2. Capture full text
                text = page.extract_text()
                if text:
                    raw_text_pages.append(text)
        
        return {
            "document_hash": file_hash,
            "fragments": fragments,
            "raw_text": "\n".join(raw_text_pages),
            "source_file": file_path
        }

    def _is_likely_transaction_table(self, table: List[List[str]]) -> bool:
        if not table or len(table) < 2: return False
        
        # Check for keywords in first 2 rows
        headers = [str(cell).lower() for row in table[:2] for cell in row if cell]
        keywords = {'date', 'amount', 'description', 'debit', 'credit', 'balance'}
        if any(k in h for h in headers for k in keywords):
            return True
            
        # Regex check for dates in rows
        date_pattern = re.compile(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}')
        for row in table[:5]:
            if any(date_pattern.search(str(cell)) for cell in row if cell):
                return True
        return False



class CSVParser(BaseParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        df = pd.read_csv(file_path)
        file_hash = self.get_file_hash(file_path)
        
        # Convert CSV rows to "fragments" of type table_row
        fragments = []
        
        # Prepend headers as first row for transformer to build column map
        fragments.append({
            "type": "table_row",
            "data": df.columns.tolist(),
            "page_number": 1
        })
        
        for _, row in df.iterrows():
            fragments.append({
                "type": "table_row",
                "data": row.tolist(),
                "page_number": 1
            })
            
        return {
            "document_hash": file_hash,
            "fragments": fragments,
            "raw_text": "", # CSV has no "raw text" usually
            "source_file": file_path
        }

class TextParser(BaseParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        file_hash = self.get_file_hash(file_path)
        return {
            "document_hash": file_hash,
            "fragments": [], # Heuristic will pick up from raw_text
            "raw_text": text,
            "source_file": file_path
        }

class DocxParser(BaseParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        import docx
        doc = docx.Document(file_path)
        fragments = []
        raw_text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                raw_text_parts.append(para.text)
                
        for table in doc.tables:
            for row in table.rows:
                fragments.append({
                    "type": "table_row",
                    "data": [cell.text for cell in row.cells],
                    "page_number": 1
                })
                
        file_hash = self.get_file_hash(file_path)
        return {
            "document_hash": file_hash,
            "fragments": fragments,
            "raw_text": "\n".join(raw_text_parts),
            "source_file": file_path
        }

class ParserFactory:
    @staticmethod
    def get_parser(file_type: str) -> BaseParser:
        ft = file_type.lower()
        if ft == 'pdf':
            return PDFParser()
        elif ft == 'csv':
            return CSVParser()
        elif ft == 'txt':
            return TextParser()
        elif ft == 'docx':
            return DocxParser()
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
